#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_builder — Geração do documento Word (.docx)

Lê um arquivo Markdown (.md) com estrutura de currículo e gera um documento
Word com layout limpo e profissional, idêntico ao modelo de referência.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


# ═══════════════════════════════════════════════════════════════════════════
#  CONSTANTES DE ESTILO
# ═══════════════════════════════════════════════════════════════════════════
FONT_NAME = 'Calibri'
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
COLOR_MEDIUM_GRAY = RGBColor(0x55, 0x55, 0x55)

NAME_SIZE = Pt(18)
SUBTITLE_SIZE = Pt(11)
CONTACT_SIZE = Pt(9)
SECTION_HEADING_SIZE = Pt(11)
JOB_TITLE_SIZE = Pt(10.5)
DATE_SIZE = Pt(9.5)
BODY_SIZE = Pt(10)
SKILL_SIZE = Pt(10)

MARGIN_TOP = Cm(1.27)
MARGIN_BOTTOM = Cm(1.27)
MARGIN_LEFT = Cm(1.91)
MARGIN_RIGHT = Cm(1.91)

BULLET_INDENT = Cm(0.63)
BULLET_HANGING = Cm(-0.32)


# ═══════════════════════════════════════════════════════════════════════════
#  FUNÇÕES AUXILIARES
# ═══════════════════════════════════════════════════════════════════════════
def _add_bottom_border(paragraph, color='000000', sz='6', space='1'):
    """Adiciona uma linha horizontal (borda inferior) ao parágrafo."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _run(paragraph, text, size=BODY_SIZE, bold=False, italic=False,
         color=COLOR_BLACK, font=FONT_NAME):
    """Cria um run com formatação e o adiciona ao parágrafo."""
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = size
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _format_inline(paragraph, text, size=BODY_SIZE, color=COLOR_BLACK):
    """Processa **negrito** inline e cria runs adequados."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            _run(paragraph, part[2:-2], size=size, bold=True, color=color)
        else:
            _run(paragraph, part, size=size, color=color)


def _para(doc, space_before=0, space_after=0, line_spacing=1.0,
          alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Cria um parágrafo vazio com formatação de espaçamento."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    return p


# ═══════════════════════════════════════════════════════════════════════════
#  BLOCOS DO CURRÍCULO
# ═══════════════════════════════════════════════════════════════════════════
def add_name(doc, name):
    p = _para(doc, space_after=0)
    _run(p, name.upper(), size=NAME_SIZE, bold=True)
    return p


def add_subtitle(doc, text):
    p = _para(doc, space_before=1, space_after=2)
    _run(p, text, size=SUBTITLE_SIZE, color=COLOR_DARK_GRAY)
    return p


def add_contact(doc, text):
    p = _para(doc, space_before=1, space_after=4)
    _run(p, text, size=CONTACT_SIZE, color=COLOR_DARK_GRAY)
    return p


def add_section_heading(doc, text):
    p = _para(doc, space_before=10, space_after=4)
    _run(p, text.upper(), size=SECTION_HEADING_SIZE, bold=True)
    _add_bottom_border(p)
    return p


def add_experience_entry(doc, title_company, date_location):
    p = _para(doc, space_before=6, space_after=0)
    _run(p, title_company, size=JOB_TITLE_SIZE, bold=True)
    if date_location:
        p2 = _para(doc, space_before=0, space_after=2)
        _run(p2, date_location, size=DATE_SIZE, italic=True,
             color=COLOR_MEDIUM_GRAY)
    return p


def add_education_entry(doc, title, years):
    p = _para(doc, space_before=4, space_after=0)
    _run(p, title, size=BODY_SIZE, bold=True)
    if years:
        p2 = _para(doc, space_before=0, space_after=2)
        _run(p2, years, size=DATE_SIZE, italic=True, color=COLOR_MEDIUM_GRAY)
    return p


def add_bullet(doc, text):
    p = _para(doc, space_before=0.5, space_after=0.5, line_spacing=1.08)
    p.paragraph_format.left_indent = BULLET_INDENT
    p.paragraph_format.first_line_indent = BULLET_HANGING
    _run(p, '• ', size=BODY_SIZE)
    _format_inline(p, text, size=BODY_SIZE)
    return p


def add_skill_line(doc, category, values):
    p = _para(doc, space_before=0.5, space_after=0.5, line_spacing=1.08)
    _run(p, category + ':', size=SKILL_SIZE, bold=True)
    _run(p, ' ' + values.strip(), size=SKILL_SIZE)
    return p


def add_language_line(doc, language, level):
    p = _para(doc, space_before=0.5, space_after=0.5, line_spacing=1.08)
    _run(p, language + ':', size=BODY_SIZE, bold=True)
    _run(p, ' ' + level.strip(), size=BODY_SIZE)
    return p


def add_body_text(doc, text):
    p = _para(doc, space_before=0, space_after=2, line_spacing=1.15)
    _format_inline(p, text, size=BODY_SIZE)
    return p


# ═══════════════════════════════════════════════════════════════════════════
#  PARSER / GERADOR PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════
def create_resume(input_file, output_file):
    """Lê o .md e gera o .docx com layout profissional."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = COLOR_BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    current_section = None
    header_done = False
    contact_lines = []
    i = 0

    while i < len(lines):
        raw = lines[i].rstrip()
        s = raw.strip()

        if not s:
            if contact_lines:
                add_contact(doc, ' '.join(contact_lines))
                contact_lines = []
            i += 1
            continue

        # ── # Nome ───────────────────────────────────────────────────────
        if s.startswith('# ') and not s.startswith('## '):
            add_name(doc, s[2:].strip())
            i += 1
            continue

        # ── ## Seção / Subtítulo ─────────────────────────────────────────
        if s.startswith('## ') and not s.startswith('### '):
            heading_text = s[3:].strip()

            if not header_done:
                add_subtitle(doc, heading_text)
                current_section = '__HEADER__'
                i += 1
                continue

            if contact_lines:
                add_contact(doc, ' '.join(contact_lines))
                contact_lines = []

            add_section_heading(doc, heading_text)
            current_section = heading_text.upper()
            i += 1
            continue

        # ── ### Subseção (experiência / educação) ────────────────────────
        if s.startswith('### '):
            title_text = s[4:].strip()
            date_location = ''

            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if nxt and not nxt.startswith('#') and not nxt.startswith('-') \
                        and not nxt.startswith('•') and not nxt.startswith('*'):
                    date_location = nxt
                    i += 1

            if current_section and ('EDUCATION' in current_section
                                    or 'EDUCAÇÃO' in current_section
                                    or 'FORMAÇÃO' in current_section):
                add_education_entry(doc, title_text, date_location)
            else:
                add_experience_entry(doc, title_text, date_location)

            i += 1
            continue

        # ── Contato (logo após subtítulo) ────────────────────────────────
        if current_section == '__HEADER__' and not header_done:
            contact_lines.append(s)
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if nxt and not nxt.startswith('#') and not nxt.startswith('-'):
                    if '•' in nxt or '@' in nxt or 'linkedin' in nxt.lower() \
                            or 'github' in nxt.lower() or 'http' in nxt.lower():
                        contact_lines.append(nxt)
                        i += 1
            add_contact(doc, ' '.join(contact_lines))
            contact_lines = []
            header_done = True
            i += 1
            continue

        # ── Bullets ──────────────────────────────────────────────────────
        if s.startswith('- ') or s.startswith('• ') or s.startswith('* '):
            text = re.sub(r'^[-•*]\s+', '', s)
            add_bullet(doc, text)
            i += 1
            continue

        # ── Skills ───────────────────────────────────────────────────────
        if current_section and 'SKILL' in current_section and ':' in s:
            cat, vals = s.split(':', 1)
            add_skill_line(doc, cat.strip(), vals)
            i += 1
            continue

        # ── Languages ────────────────────────────────────────────────────
        if current_section and ('LANGUAGE' in current_section
                                or 'IDIOMA' in current_section) and ':' in s:
            lang, lvl = s.split(':', 1)
            add_language_line(doc, lang.strip(), lvl)
            i += 1
            continue

        # ── Texto normal ─────────────────────────────────────────────────
        if s:
            add_body_text(doc, s)

        i += 1

    if contact_lines:
        add_contact(doc, ' '.join(contact_lines))

    doc.save(output_file)
    print(f'Curriculo gerado com sucesso: {output_file}')


def create_resume_from_text(md_text, output_file):
    """Gera .docx a partir de texto markdown (string) em vez de arquivo."""
    import tempfile, os
    tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.md',
                                      delete=False, encoding='utf-8')
    tmp.write(md_text)
    tmp.close()
    try:
        create_resume(tmp.name, output_file)
    finally:
        os.unlink(tmp.name)
